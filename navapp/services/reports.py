from __future__ import annotations

import hashlib
import os
import posixpath
import re
import shutil
import subprocess
import tempfile
from datetime import datetime
from pathlib import Path
from xml.etree import ElementTree
from zipfile import ZipFile

import matplotlib

matplotlib.use("Agg")
import matplotlib.dates as mdates
import matplotlib.pyplot as plt
from django.conf import settings
from django.core.exceptions import ObjectDoesNotExist, ValidationError
from django.db import transaction
from django.utils import timezone
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from docxtpl import DocxTemplate, InlineImage

from navapp.models import (
    AuditLog,
    GeneratedFile,
    QuarterlyReport,
)
from navapp.services.calculations import CalculationValidationError, calculate_for_report
from navapp.services.commentary import parse_commentary


class ReportGenerationError(RuntimeError):
    pass


REQUIRED_CUSTOM_PLACEHOLDERS = {
    "fund_name",
    "report_quarter",
    "report_date",
    "quarterly_rows",
    "nav_chart",
    "manager_commentary",
    "disclaimer",
}


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def external_excel_relationships(path: Path) -> list[str]:
    matches: list[str] = []
    with ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.endswith(".rels"):
                continue
            text = archive.read(name).decode("utf-8", errors="ignore")
            for relation in re.findall(r"<Relationship\b[^>]+>", text):
                lower = relation.lower()
                if 'targetmode="external"' in lower and (
                    ".xlsx" in lower or ".xls" in lower or "oleobject" in lower or "excel" in lower
                ):
                    matches.append(f"{name}: {relation}")
    return matches


def audit_docx_package(path: Path) -> dict[str, object]:
    """Return a JSON-serializable structural audit for a generated DOCX package."""
    with ZipFile(path) as archive:
        names = set(archive.namelist())
        corrupt_member = archive.testzip()
        external_relationships: list[str] = []
        missing_relationship_targets: list[str] = []
        for relationship_name in sorted(name for name in names if name.endswith(".rels")):
            try:
                root = ElementTree.fromstring(archive.read(relationship_name))
            except ElementTree.ParseError:
                missing_relationship_targets.append(f"Malformed relationships: {relationship_name}")
                continue
            base = (
                ""
                if relationship_name == "_rels/.rels"
                else str(Path(relationship_name).parent.parent).replace("\\", "/")
            )
            for relation in root:
                target = relation.attrib.get("Target", "")
                if relation.attrib.get("TargetMode", "").lower() == "external":
                    external_relationships.append(f"{relationship_name}: {target}")
                    continue
                if not target or target.startswith("#"):
                    continue
                resolved = posixpath.normpath(
                    target.lstrip("/") if target.startswith("/") else posixpath.join(base, target)
                )
                if resolved not in names:
                    missing_relationship_targets.append(
                        f"{relationship_name}: {target} -> {resolved}"
                    )
        document_xml = archive.read("word/document.xml")
        root = ElementTree.fromstring(document_xml)
        namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        table_widths = [
            int(node.attrib.get(qn("w:w"), "0"))
            for node in root.findall(".//w:tbl/w:tblPr/w:tblW", namespaces)
        ]
        fixed_row_heights = [
            node.attrib.get(qn("w:val"), "") for node in root.findall(".//w:trHeight", namespaces)
        ]
        media = sorted(name for name in names if name.startswith("word/media/"))
        spreadsheets = sorted(
            name
            for name in names
            if name.startswith("word/embeddings/")
            or name.lower().endswith((".xls", ".xlsx", ".xlsm"))
        )
        footer_texts = []
        for name in sorted(
            item for item in names if item.startswith("word/footer") and item.endswith(".xml")
        ):
            footer_root = ElementTree.fromstring(archive.read(name))
            footer_texts.append(
                "".join(node.text or "" for node in footer_root.findall(".//w:t", namespaces))
            )
        drawing_count = len(root.findall(".//w:drawing", namespaces))

    document = Document(path)
    metadata = {
        "title": document.core_properties.title or "",
        "subject": document.core_properties.subject or "",
        "author": document.core_properties.author or "",
        "keywords": document.core_properties.keywords or "",
        "comments": document.core_properties.comments or "",
    }
    valid = all(
        (
            corrupt_member is None,
            not external_relationships,
            not missing_relationship_targets,
            not spreadsheets,
            bool(media),
            drawing_count >= 1,
            bool(table_widths) and all(width == 9864 for width in table_widths),
            not fixed_row_heights,
            bool(footer_texts)
            and all("Report " in text and "Generated " in text for text in footer_texts),
            bool(metadata["title"]),
            "formula=legacy_excel_v1" in metadata["keywords"],
        )
    )
    return {
        "valid": valid,
        "path": str(path.resolve()),
        "size": path.stat().st_size,
        "sha256": sha256_file(path),
        "corrupt_member": corrupt_member,
        "external_relationships": external_relationships,
        "external_excel_relationships": external_excel_relationships(path),
        "missing_relationship_targets": missing_relationship_targets,
        "embedded_media": media,
        "embedded_spreadsheets": spreadsheets,
        "drawing_count": drawing_count,
        "table_widths_dxa": table_widths,
        "fixed_row_heights": fixed_row_heights,
        "footer_texts": footer_texts,
        "metadata": metadata,
    }


def validate_custom_template(path: Path) -> None:
    if path.suffix.lower() != ".docx" or not path.is_file():
        raise ReportGenerationError("自訂範本必須是可讀取的 .docx 檔案。")
    try:
        with ZipFile(path) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
            package_text = "".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if name.endswith(".xml")
            )
    except (KeyError, OSError) as exc:
        raise ReportGenerationError(f"自訂 DOCX 封裝無效：{exc}") from exc
    missing = [name for name in REQUIRED_CUSTOM_PLACEHOLDERS if name not in package_text]
    if missing:
        raise ReportGenerationError("自訂範本缺少必要的預留位置：" + ", ".join(sorted(missing)))
    if "{{" not in document_xml and "{%" not in document_xml:
        raise ReportGenerationError("自訂範本不包含 docxtpl 預留位置。")
    external = external_excel_relationships(path)
    if external:
        raise ReportGenerationError("自訂範本包含外部 Excel 關聯。")


def build_current_snapshot(report: QuarterlyReport) -> dict[str, object]:
    calculation = calculate_for_report(report)
    fund = report.fund
    share_class = report.share_class
    resolved = fund.resolved()
    rfr: dict[str, object] | None = None
    try:
        source = report.rfr_snapshot
        rfr = {
            "provider": source.provider,
            "series": source.series,
            "annual_value_decimal": str(source.annual_value_decimal),
            "is_manual": source.is_manual,
            "override_reason": source.override_reason,
            "raw_checksum": source.raw_checksum,
            "observations": [
                {
                    "position": item.position,
                    "date": item.observation_date.isoformat(),
                    "value_percent": str(item.value_percent),
                }
                for item in source.observations.all()
            ],
        }
    except ObjectDoesNotExist:
        pass
    date_statement = str(resolved["date_statement_template"]).format(
        report_date=report.report_date.strftime("%-d %B %Y")
        if os.name != "nt"
        else report.report_date.strftime("%d %B %Y").lstrip("0")
    )
    return {
        "identity": {
            "fund_id": fund.pk,
            "fund_legal_name": fund.legal_name,
            "fund_display_name": fund.display_name,
            "fund_short_code": fund.short_code,
            "share_class_id": share_class.pk,
            "share_class_name": share_class.name,
            "share_class_code": share_class.code,
            "report_id": report.pk,
            "year": report.year,
            "quarter": report.quarter,
            "version": report.version,
            "report_date": report.report_date.isoformat(),
        },
        "fund": {
            "structure": fund.structure,
            "domicile": fund.domicile,
            "investment_objective": fund.investment_objective,
            "performance_note": fund.performance_note,
            "professional_statement": resolved["professional_statement"],
            "date_statement": date_statement,
            "brand_colour": resolved["brand_colour"],
            "logo_path": resolved["logo_path"],
            "disclaimer": resolved["disclaimer"],
            "disclaimer_version": resolved["disclaimer_version"],
            "strategies": [item.text for item in fund.strategy_highlights.all()],
            "parties": [
                {"label": item.display_label, "value": item.value} for item in fund.parties.all()
            ],
            "terms": [
                {"label": item.display_label, "value": item.value_text}
                for item in fund.terms.filter(display_in_report=True)
            ],
            "contacts": [
                {
                    "role": item.role,
                    "name": item.name,
                    "email": item.email,
                    "phone": item.phone,
                    "address": item.address,
                }
                for item in fund.contacts.filter(display_in_report=True)
            ],
        },
        "share_class": {
            "inception_date": share_class.inception_date.isoformat(),
            "inception_nav": str(share_class.inception_nav),
            "currency": share_class.currency,
            "return_basis": share_class.return_basis,
        },
        "calculation": calculation,
        "rfr": rfr,
        "commentary": {
            "title": report.commentary_title,
            "markdown": report.commentary_markdown,
            "author": report.commentary_author,
            "date": report.commentary_date.isoformat() if report.commentary_date else None,
        },
        "formula_version": report.formula_version,
        "captured_at": timezone.now().isoformat(),
    }


def generate_nav_chart(snapshot: dict[str, object], output_path: Path) -> None:
    monthly = snapshot["calculation"]["monthly"]
    dates = [datetime.fromisoformat(row["valuation_month"]) for row in monthly]
    values = [float(row["nav"]) for row in monthly]
    inception = snapshot["share_class"]
    inception_date = datetime.fromisoformat(inception["inception_date"])
    if not dates or inception_date < dates[0]:
        dates.insert(0, inception_date)
        values.insert(0, float(inception["inception_nav"]))
    colour = snapshot["fund"]["brand_colour"] or "#183B73"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    fig, axis = plt.subplots(figsize=(9.6, 4.2), dpi=180)
    try:
        fig.patch.set_facecolor("white")
        axis.set_facecolor("white")
        axis.plot(dates, values, color=colour, linewidth=2.4, marker="o", markersize=3.4)
        axis.set_ylabel("NAV per Share")
        axis.grid(axis="y", color="#DCE3EA", linewidth=0.7)
        axis.spines[["top", "right"]].set_visible(False)
        axis.spines[["left", "bottom"]].set_color("#9AA8B6")
        interval = 3 if len(dates) <= 36 else 6
        axis.xaxis.set_major_locator(mdates.MonthLocator(interval=interval))
        axis.xaxis.set_major_formatter(mdates.DateFormatter("%b %Y"))
        fig.autofmt_xdate(rotation=35, ha="right")
        axis.margins(x=0.02)
        fig.tight_layout()
        fig.savefig(output_path, bbox_inches="tight", facecolor="white")
    finally:
        plt.close(fig)


def _set_run_font(run, name: str = "Arial", size: float | None = None, **kwargs) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    for key, value in kwargs.items():
        if key == "color":
            run.font.color.rgb = value
        else:
            setattr(run.font, key, value)


def _configure_styles(document: Document, colour: RGBColor) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    heading_tokens = {
        "Heading 1": (16, 16, 8),
        "Heading 2": (13, 12, 6),
        "Heading 3": (12, 8, 4),
    }
    for name, (size, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = colour
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    if "Disclaimer" not in styles:
        disclaimer = styles.add_style("Disclaimer", WD_STYLE_TYPE.PARAGRAPH)
    else:
        disclaimer = styles["Disclaimer"]
    disclaimer.font.name = "Arial"
    disclaimer.font.size = Pt(10)
    disclaimer.paragraph_format.space_after = Pt(6)
    disclaimer.paragraph_format.line_spacing = 1.15


def _set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_table_geometry(table, widths: list[int], total_width: int = 9864) -> None:
    if sum(widths) != total_width:
        raise ReportGenerationError("DOCX 表格欄寬總和必須等於內容寬度。")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("tblW", "tblInd", "tblLayout"):
        existing = tbl_pr.find(qn(f"w:{tag}"))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(total_width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell, width in zip(row.cells, widths, strict=True):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def _style_table_text(table, header_fill: str = "E8EEF5") -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                _shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    _set_run_font(run, size=9.5)
                    run.bold = row_index == 0


def _write_footer(paragraph, snapshot: dict[str, object]) -> None:
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    identity = snapshot["identity"]
    generated_at = datetime.fromisoformat(str(snapshot["captured_at"])).strftime(
        "%Y-%m-%d %H:%M UTC"
    )
    run = paragraph.add_run(
        f"{identity['fund_short_code']} | Report {identity['report_id']} | "
        f"{identity['year']} Q{identity['quarter']} v{identity['version']} | "
        f"Generated {generated_at} | Page "
    )
    _set_run_font(run, size=8, color=RGBColor(100, 112, 125))
    field_run = paragraph.add_run()
    _set_run_font(field_run, size=8, color=RGBColor(100, 112, 125))
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    field_run._r.extend([field_begin, instr, field_end])


def _add_footer(document: Document, snapshot: dict[str, object]) -> None:
    document.settings.odd_and_even_pages_header_footer = False
    section = document.sections[0]
    section.different_first_page_header_footer = False
    for footer in (section.footer, section.even_page_footer, section.first_page_footer):
        _write_footer(footer.paragraphs[0], snapshot)


def _add_commentary(document: Document, markdown_value: str) -> list:
    paragraphs = []
    for block in parse_commentary(markdown_value):
        style = {
            "bullet": "List Bullet",
            "number": "List Number",
            "paragraph": None,
        }[block.kind]
        paragraph = document.add_paragraph(style=style)
        paragraphs.append(paragraph)
        for segment in block.segments:
            run = paragraph.add_run(segment.text)
            run.bold = segment.bold
            run.italic = segment.italic
    return paragraphs


def build_builtin_docx(snapshot: dict[str, object], chart_path: Path, output_path: Path) -> None:
    document = Document()
    identity = snapshot["identity"]
    core = document.core_properties
    core.title = (
        f"{identity['fund_display_name']} - {identity['year']} "
        f"Q{identity['quarter']} Quarterly Report"
    )
    core.subject = f"NAV quarterly report for {identity['share_class_name']}"
    core.author = "NAV Quarterly Reporting"
    core.keywords = (
        f"report_id={identity['report_id']}; version={identity['version']}; "
        f"formula={snapshot['formula_version']}"
    )
    core.comments = f"Snapshot captured at {snapshot['captured_at']}"
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.right_margin = Mm(18)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.header_distance = Mm(10)
    section.footer_distance = Mm(10)
    colour_hex = str(snapshot["fund"]["brand_colour"] or "#183B73").lstrip("#")
    colour = RGBColor.from_string(colour_hex.upper())
    _configure_styles(document, colour)
    _add_footer(document, snapshot)

    fund = snapshot["fund"]
    share = snapshot["share_class"]
    logo_path = Path(str(fund["logo_path"])) if fund["logo_path"] else None
    if logo_path and logo_path.is_file():
        logo = document.add_paragraph()
        logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        logo.paragraph_format.space_after = Pt(3)
        logo.add_run().add_picture(str(logo_path), height=Mm(13))
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(identity["fund_display_name"])
    _set_run_font(run, size=25, bold=True, color=colour)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle.paragraph_format.keep_with_next = True
    run = subtitle.add_run(f"{identity['year']} Q{identity['quarter']}")
    _set_run_font(run, size=15, bold=True, color=RGBColor(55, 70, 85))
    for line in (fund["professional_statement"], fund["date_statement"]):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(1)
        run = paragraph.add_run(str(line))
        _set_run_font(run, size=9.5, color=RGBColor(75, 85, 95))
    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:color"), colour_hex.upper())
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    document.add_heading("Investment Objective", level=1)
    document.add_paragraph(str(fund["investment_objective"]))
    document.add_heading("Strategy Highlights and Characteristics", level=1)
    for strategy in fund["strategies"]:
        document.add_paragraph(str(strategy), style="List Bullet")

    document.add_heading("Fund Performance (Net Quarterly Returns)", level=1)
    matrix = snapshot["calculation"]["quarterly_matrix"]
    performance = document.add_table(rows=1, cols=6)
    performance.style = "Table Grid"
    headers = ["Year", "Q1", "Q2", "Q3", "Q4", "YTD"]
    for cell, label in zip(performance.rows[0].cells, headers, strict=True):
        cell.text = label
    for year, values in sorted(matrix.items()):
        cells = performance.add_row().cells
        row_values = [year] + [values[key]["display"] for key in ("q1", "q2", "q3", "q4", "ytd")]
        for cell, value in zip(cells, row_values, strict=True):
            cell.text = str(value)
    _set_table_geometry(performance, [1120, 1748, 1748, 1748, 1748, 1752])
    _repeat_header(performance.rows[0])
    _style_table_text(performance, "394B59")
    for cell in performance.rows[0].cells:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    if fund["performance_note"]:
        note = document.add_paragraph(str(fund["performance_note"]))
        note.paragraph_format.space_before = Pt(4)
        for run in note.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(105, 110, 115)
            run.font.size = Pt(9)

    document.add_heading("Fund Performance (Graph)", level=1)
    document.add_picture(str(chart_path), width=Inches(6.55))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Fund Statistics", level=1)
    metric_labels = [
        ("Inception Date", share["inception_date"]),
        ("Return ITD", snapshot["calculation"]["metrics_display"]["itd_return"]),
        ("Return YTD", snapshot["calculation"]["metrics_display"]["ytd_return"]),
        ("Annualized Return", snapshot["calculation"]["metrics_display"]["annualized_return"]),
        ("Positive Months", snapshot["calculation"]["metrics_display"]["positive_months"]),
        ("Negative Months", snapshot["calculation"]["metrics_display"]["negative_months"]),
        (
            "Annualized Volatility",
            snapshot["calculation"]["metrics_display"]["annualized_volatility"],
        ),
        ("Maximum Drawdown", snapshot["calculation"]["metrics_display"]["maximum_drawdown"]),
        ("Sharpe Ratio", snapshot["calculation"]["metrics_display"]["sharpe_ratio"]),
    ]
    stats = document.add_table(rows=1, cols=2)
    stats.style = "Table Grid"
    stats.rows[0].cells[0].text = "Statistic"
    stats.rows[0].cells[1].text = "Value"
    for label, value in metric_labels:
        row = stats.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)
    _set_table_geometry(stats, [6000, 3864])
    _repeat_header(stats.rows[0])
    _style_table_text(stats)

    document.add_heading("Manager Commentary", level=1)
    if snapshot["commentary"]["title"]:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(str(snapshot["commentary"]["title"]))
        run.bold = True
    commentary_paragraphs = _add_commentary(document, str(snapshot["commentary"]["markdown"]))
    if snapshot["commentary"]["author"]:
        if commentary_paragraphs:
            commentary_paragraphs[-1].paragraph_format.keep_with_next = True
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run(f"— {snapshot['commentary']['author']}")
        run.italic = True

    document.add_page_break()
    document.add_heading("General Information", level=1)
    general_rows = [
        ("Structure", fund["structure"]),
        ("Domicile", fund["domicile"]),
        ("Currency", share["currency"]),
        ("Return Basis", share["return_basis"].title()),
        *[(row["label"], row["value"]) for row in fund["parties"]],
        *[(row["label"], row["value"]) for row in fund["terms"]],
    ]
    general = document.add_table(rows=1, cols=2)
    general.style = "Table Grid"
    general.rows[0].cells[0].text = "Item"
    general.rows[0].cells[1].text = "Details"
    for label, value in general_rows:
        row = general.add_row().cells
        row[0].text = str(label)
        row[1].text = str(value)
    _set_table_geometry(general, [3000, 6864])
    _repeat_header(general.rows[0])
    _style_table_text(general)

    document.add_heading("Contacts", level=1)
    contacts = document.add_table(rows=1, cols=2)
    contacts.style = "Table Grid"
    contacts.rows[0].cells[0].text = "Contact"
    contacts.rows[0].cells[1].text = "Details"
    for contact in fund["contacts"]:
        details = [contact["name"], contact["email"], contact["phone"], contact["address"]]
        row = contacts.add_row().cells
        row[0].text = str(contact["role"])
        row[1].text = "\n".join(str(value) for value in details if value)
    _set_table_geometry(contacts, [3000, 6864])
    _repeat_header(contacts.rows[0])
    _style_table_text(contacts)

    document.add_page_break()
    document.add_heading("Disclaimer", level=1)
    for paragraph_text in str(fund["disclaimer"]).replace("\r\n", "\n").split("\n\n"):
        if paragraph_text.strip():
            document.add_paragraph(paragraph_text.strip(), style="Disclaimer")
    provenance = document.add_paragraph(style="Disclaimer")
    rfr = snapshot.get("rfr") or {}
    provenance.add_run("Calculation and provenance: ").bold = True
    provenance.add_run(
        f"{snapshot['formula_version']} | Report ID {identity['report_id']} | "
        f"Version {identity['version']} | Snapshot {snapshot['captured_at']}"
        + (f" | RFR {rfr.get('provider')} / {rfr.get('series')}" if rfr else "")
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def build_custom_docx(
    snapshot: dict[str, object], chart_path: Path, template_path: Path, output_path: Path
) -> None:
    validate_custom_template(template_path)
    template = DocxTemplate(str(template_path))
    matrix = snapshot["calculation"]["quarterly_matrix"]
    context = {
        "fund_name": snapshot["identity"]["fund_display_name"],
        "share_class": snapshot["identity"]["share_class_name"],
        "report_quarter": f"{snapshot['identity']['year']} Q{snapshot['identity']['quarter']}",
        "report_date": snapshot["identity"]["report_date"],
        "investment_objective": snapshot["fund"]["investment_objective"],
        "strategy_highlights": snapshot["fund"]["strategies"],
        "quarterly_rows": [
            {"year": year, **{key: value["display"] for key, value in row.items()}}
            for year, row in sorted(matrix.items())
        ],
        "nav_chart": InlineImage(template, str(chart_path), width=Mm(165)),
        "fund_statistics": snapshot["calculation"]["metrics_display"],
        "manager_commentary": snapshot["commentary"]["markdown"],
        "general_information": [*snapshot["fund"]["parties"], *snapshot["fund"]["terms"]],
        "contacts": snapshot["fund"]["contacts"],
        "disclaimer": snapshot["fund"]["disclaimer"],
    }
    template.render(context, autoescape=True)
    template.save(str(output_path))


def convert_docx_to_pdf(docx_path: Path, output_dir: Path) -> Path:
    binary = shutil.which(settings.LIBREOFFICE_BINARY)
    if not binary:
        raise ReportGenerationError(f"找不到 LibreOffice 執行檔 {settings.LIBREOFFICE_BINARY!r}。")
    output_dir.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="nav-lo-profile-") as profile:
        profile_uri = Path(profile).resolve().as_uri()
        command = [
            binary,
            "--headless",
            f"-env:UserInstallation={profile_uri}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir),
            str(docx_path),
        ]
        try:
            result = subprocess.run(
                command,
                check=False,
                capture_output=True,
                text=True,
                timeout=settings.REPORT_CONVERSION_TIMEOUT,
            )
        except subprocess.TimeoutExpired as exc:
            raise ReportGenerationError("LibreOffice PDF 轉換逾時。") from exc
    pdf_path = output_dir / f"{docx_path.stem}.pdf"
    if result.returncode != 0 or not pdf_path.exists() or pdf_path.stat().st_size == 0:
        raise ReportGenerationError(
            "LibreOffice PDF 轉換失敗："
            f"stdout={result.stdout[-1000:]} stderr={result.stderr[-1000:]}"
        )
    return pdf_path


def _save_generated_file(report: QuarterlyReport, file_type: str, path: Path) -> GeneratedFile:
    media_root = Path(settings.MEDIA_ROOT).resolve()
    relative = path.resolve().relative_to(media_root).as_posix()
    obj, _ = GeneratedFile.objects.update_or_create(
        report=report,
        file_type=file_type,
        defaults={
            "storage_path": relative,
            "sha256": sha256_file(path),
            "size": path.stat().st_size,
        },
    )
    return obj


def generate_report_files(report: QuarterlyReport, actor=None) -> list[GeneratedFile]:
    if report.status in {QuarterlyReport.Status.FINAL, QuarterlyReport.Status.STALE}:
        raise ReportGenerationError("已定稿報告的檔案不可修改；請建立新版本。")
    try:
        snapshot = build_current_snapshot(report)
        report.snapshot = snapshot
        report.generation_error = ""
        report.save(update_fields=["snapshot", "generation_error", "updated_at"])
        output_dir = Path(settings.MEDIA_ROOT) / "reports" / str(report.pk) / f"v{report.version}"
        output_dir.mkdir(parents=True, exist_ok=True)
        chart_path = output_dir / "nav-chart.png"
        docx_path = output_dir / "quarterly-report.docx"
        generate_nav_chart(snapshot, chart_path)
        if report.fund.custom_docx_template:
            build_custom_docx(
                snapshot,
                chart_path,
                Path(report.fund.custom_docx_template.path),
                docx_path,
            )
        else:
            build_builtin_docx(snapshot, chart_path, docx_path)
        if external_excel_relationships(docx_path):
            raise ReportGenerationError("產生的 DOCX 包含外部 Excel 關聯。")
        with transaction.atomic():
            docx_file = _save_generated_file(report, GeneratedFile.FileType.DOCX, docx_path)
            AuditLog.objects.create(
                actor=actor,
                entity_type="QuarterlyReport",
                entity_id=str(report.pk),
                action="GENERATE_DOCX",
                after_json={"docx": docx_file.sha256},
            )
        pdf_path = convert_docx_to_pdf(docx_path, output_dir)
        with transaction.atomic():
            pdf_file = _save_generated_file(report, GeneratedFile.FileType.PDF, pdf_path)
            report.status = QuarterlyReport.Status.READY
            report.save(update_fields=["status", "updated_at"])
            AuditLog.objects.create(
                actor=actor,
                entity_type="QuarterlyReport",
                entity_id=str(report.pk),
                action="GENERATE",
                after_json={"docx": docx_file.sha256, "pdf": pdf_file.sha256},
            )
        return [docx_file, pdf_file]
    except Exception as exc:
        report.generation_error = str(exc)
        report.status = QuarterlyReport.Status.GENERATION_FAILED
        report.save(update_fields=["generation_error", "status", "updated_at"])
        if isinstance(exc, (ReportGenerationError, CalculationValidationError)):
            raise
        raise ReportGenerationError(str(exc)) from exc


def finalization_issues(report: QuarterlyReport) -> list[str]:
    issues: list[str] = []
    required_values = {
        "基金法定名稱": report.fund.legal_name,
        "基金顯示名稱": report.fund.display_name,
        "基金架構": report.fund.structure,
        "基金註冊地": report.fund.domicile,
        "投資目標": report.fund.investment_objective,
        "股份類別名稱": report.share_class.name,
        "股份類別貨幣": report.share_class.currency,
    }
    issues.extend(f"必須填寫{label}。" for label, value in required_values.items() if not value)
    if not report.commentary_markdown.strip():
        issues.append("必須填寫基金經理評論。")
    try:
        calculate_for_report(report)
    except CalculationValidationError as exc:
        issues.extend(exc.issues)
    try:
        rfr = report.rfr_snapshot
        if not rfr.is_manual:
            observations = list(rfr.observations.all())
            if len(observations) != 12:
                issues.append("必須有正好 12 筆無風險利率觀察值。")
            elif (
                observations[-1].observation_date.year != report.report_date.year
                or observations[-1].observation_date.month != report.report_date.month
            ):
                issues.append("無風險利率截止月份必須與報告結束月份相同。")
            if any(item.observation_date > report.report_date for item in observations):
                issues.append("無風險利率觀察日期不得超過報告截止日。")
        elif not rfr.override_reason.strip():
            issues.append("手動覆寫無風險利率必須填寫原因。")
    except ObjectDoesNotExist:
        issues.append("必須建立無風險利率快照。")
    files = {item.file_type: item for item in report.files.all()}
    for file_type in (GeneratedFile.FileType.DOCX, GeneratedFile.FileType.PDF):
        item = files.get(file_type)
        if not item or not item.absolute_path.exists() or item.absolute_path.stat().st_size == 0:
            issues.append(f"必須有非空白的 {file_type} 檔案。")
        elif sha256_file(item.absolute_path) != item.sha256:
            issues.append(f"已儲存的 {file_type} 檔案雜湊值與檔案不符。")
    if not report.snapshot:
        issues.append("必須建立不可修改的報告快照。")
    return issues


@transaction.atomic
def finalize_report(report: QuarterlyReport, actor) -> QuarterlyReport:
    issues = finalization_issues(report)
    if issues:
        raise ValidationError(issues)
    report.status = QuarterlyReport.Status.FINAL
    report.finalized_at = timezone.now()
    report.finalized_by = actor
    report.save(update_fields=["status", "finalized_at", "finalized_by", "updated_at"])
    AuditLog.objects.create(
        actor=actor,
        entity_type="QuarterlyReport",
        entity_id=str(report.pk),
        action="FINALIZE",
        after_json={"status": report.status, "snapshot": report.snapshot},
    )
    return report


def mark_affected_reports_stale(nav_record, actor, reason: str) -> int:
    affected = QuarterlyReport.objects.filter(
        share_class=nav_record.share_class,
        report_date__gte=nav_record.valuation_month,
        status=QuarterlyReport.Status.FINAL,
    )
    count = affected.update(status=QuarterlyReport.Status.STALE)
    if count:
        AuditLog.objects.create(
            actor=actor,
            entity_type="NAVRecord",
            entity_id=str(nav_record.pk),
            action="MARK_REPORTS_STALE",
            after_json={"reports_marked_stale": count},
            reason=reason,
        )
    return count


def _mark_source_reports_stale(
    *,
    reports,
    actor,
    entity_type: str,
    entity_id: str,
    reason: str,
) -> int:
    count = reports.filter(status=QuarterlyReport.Status.FINAL).update(
        status=QuarterlyReport.Status.STALE
    )
    if count:
        AuditLog.objects.create(
            actor=actor,
            entity_type=entity_type,
            entity_id=entity_id,
            action="MARK_REPORTS_STALE",
            after_json={"reports_marked_stale": count},
            reason=reason,
        )
    return count


def mark_fund_reports_stale(fund, actor, reason: str) -> int:
    """Mark finalized reports stale after reportable fund data changes."""
    return _mark_source_reports_stale(
        reports=QuarterlyReport.objects.filter(fund=fund),
        actor=actor,
        entity_type="Fund",
        entity_id=str(fund.pk),
        reason=reason,
    )


def mark_share_class_reports_stale(share_class, actor, reason: str) -> int:
    """Mark finalized reports stale after reportable share-class data changes."""
    return _mark_source_reports_stale(
        reports=QuarterlyReport.objects.filter(share_class=share_class),
        actor=actor,
        entity_type="ShareClass",
        entity_id=str(share_class.pk),
        reason=reason,
    )


def mark_organization_reports_stale(organization, actor, reason: str) -> int:
    """Mark finalized reports stale after organization-level report settings change."""
    return _mark_source_reports_stale(
        reports=QuarterlyReport.objects.all(),
        actor=actor,
        entity_type="OrganizationSettings",
        entity_id=str(organization.pk),
        reason=reason,
    )
