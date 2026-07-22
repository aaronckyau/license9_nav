from __future__ import annotations

import hashlib
import shutil
from datetime import date
from decimal import Decimal
from pathlib import Path
from types import SimpleNamespace
from zipfile import ZipFile

import pytest
from django.contrib.auth import get_user_model
from django.core.exceptions import ValidationError
from django.urls import reverse
from docx import Document
from docx.oxml.ns import qn
from matplotlib.axes import Axes

from navapp.models import (
    Fund,
    GeneratedFile,
    NAVRecord,
    OrganizationSettings,
    QuarterlyReport,
    ShareClass,
)
from navapp.services import reports
from navapp.services.calculations import NavPoint, calculate_performance
from navapp.services.reports import (
    ReportGenerationError,
    audit_docx_package,
    build_boya_reference_docx,
    build_builtin_docx,
    build_current_snapshot,
    external_excel_relationships,
    finalize_report,
    generate_nav_chart,
    generate_report_files,
    mark_affected_reports_stale,
    mark_fund_reports_stale,
    mark_organization_reports_stale,
    mark_share_class_reports_stale,
    reset_affected_mutable_reports,
    sha256_file,
)
from navapp.services.rfr import set_manual_snapshot


@pytest.mark.parametrize(
    ("point_count", "expected_interval"),
    [(1, 1), (18, 1), (19, 2), (36, 2), (37, 3), (72, 3), (73, 6)],
)
def test_nav_chart_tick_interval_adapts_to_history_length(point_count, expected_interval):
    assert reports._nav_chart_tick_interval(point_count) == expected_interval


def test_nav_chart_aligns_month_end_values_with_their_month_labels(tmp_path, monkeypatch):
    snapshot = {
        "identity": {"report_language": "zh-Hant"},
        "fund": {"brand_colour": "#183B73"},
        "share_class": {"inception_date": "2025-11-24", "inception_nav": "100.00"},
        "calculation": {
            "monthly": [
                {"valuation_month": "2025-11-30", "nav": "100.00"},
                {"valuation_month": "2025-12-31", "nav": "99.72"},
                {"valuation_month": "2026-01-31", "nav": "99.80"},
                {"valuation_month": "2026-02-28", "nav": "100.85"},
                {"valuation_month": "2026-03-31", "nav": "92.39"},
                {"valuation_month": "2026-04-30", "nav": "94.54"},
                {"valuation_month": "2026-05-31", "nav": "91.39"},
            ]
        },
    }
    plotted_dates = []
    original_plot = Axes.plot

    def capture_plot(axis, dates, values, *args, **kwargs):
        plotted_dates.extend(dates)
        return original_plot(axis, dates, values, *args, **kwargs)

    monkeypatch.setattr(Axes, "plot", capture_plot)

    generate_nav_chart(snapshot, tmp_path / "may-chart.png")

    assert [value.strftime("%Y-%m-%d") for value in plotted_dates] == [
        "2025-11-01",
        "2025-12-01",
        "2026-01-01",
        "2026-02-01",
        "2026-03-01",
        "2026-04-01",
        "2026-05-01",
    ]


@pytest.fixture
def report_fixture(db, settings, tmp_path):
    settings.MEDIA_ROOT = tmp_path / "media"
    org = OrganizationSettings.load()
    org.default_disclaimer = "Organization disclaimer"
    org.save()
    user = get_user_model().objects.create_user("reporter", password="safe-password")
    fund = Fund.objects.create(
        legal_name="Example Fund LPF",
        display_name="Example Fund",
        short_code="example",
        structure="Limited Partnership Fund",
        domicile="Hong Kong",
        investment_objective="Preserve capital and compound returns.",
        performance_note="Returns are net of applicable fees.",
        use_org_professional_statement=False,
        professional_investor_statement_override="Professional investors only - fund override",
        use_org_disclaimer=False,
        disclaimer_override="Fund-specific legal disclaimer.",
        disclaimer_version_override="FUND-1",
    )
    share = ShareClass.objects.create(
        fund=fund,
        name="Class A",
        code="a",
        inception_date=date(2024, 1, 1),
        inception_nav=Decimal("100"),
        currency="USD",
    )
    records = []
    for month, nav in ((1, "101"), (2, "103"), (3, "105")):
        day = 31 if month in {1, 3} else 29
        records.append(
            NAVRecord(
                share_class=share,
                valuation_month=date(2024, month, day),
                valuation_date=date(2024, month, day),
                nav_per_share=Decimal(nav),
            )
        )
    NAVRecord.objects.bulk_create(records)
    report = QuarterlyReport.objects.create(
        fund=fund,
        share_class=share,
        year=2024,
        quarter=1,
        report_date=date(2024, 3, 31),
        commentary_title="Quarterly review",
        commentary_markdown="**Disciplined execution** remained central.\n\n- Managed risk\n- Preserved liquidity",
        commentary_author="Portfolio Manager",
        commentary_date=date(2024, 3, 31),
        created_by=user,
    )
    set_manual_snapshot(report, Decimal("4.25"), "Approved legacy reconciliation", user)
    return report, user, tmp_path


def _build_docx(report, tmp_path: Path) -> tuple[dict[str, object], Path]:
    snapshot = build_current_snapshot(report)
    chart = tmp_path / "chart.png"
    output = tmp_path / "report.docx"
    generate_nav_chart(snapshot, chart)
    build_builtin_docx(snapshot, chart, output)
    return snapshot, output


@pytest.mark.django_db
def test_docx_package_contains_sections_table_values_chart_and_no_excel_links(report_fixture):
    report, _, tmp_path = report_fixture
    snapshot, output = _build_docx(report, tmp_path)
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "投資目標" in text
    assert "基金表現（季度淨回報）" in text
    assert "基金表現（圖表）" in text
    assert "基金經理評論" in text
    assert "Fund-specific legal disclaimer" in text
    assert "Professional investors only - fund override" in text
    assert "Calculation and provenance:" not in text
    section_positions = [
        text.index(label)
        for label in (
            "投資目標",
            "策略重點及特點",
            "基金表現（季度淨回報）",
            "基金表現（圖表）",
            "基金經理評論",
            "一般資料",
            "聯絡資料",
            "免責聲明",
        )
    ]
    assert section_positions == sorted(section_positions)
    assert snapshot["calculation"]["quarterly_matrix"]["2024"]["q1"]["display"] == "5.00%"
    assert any(
        "5.00%" in cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    with ZipFile(output) as package:
        assert any(name.startswith("word/media/") for name in package.namelist())
        assert package.testzip() is None
    assert external_excel_relationships(output) == []
    package_audit = audit_docx_package(output)
    assert package_audit["valid"] is True
    assert package_audit["external_relationships"] == []
    assert package_audit["missing_relationship_targets"] == []
    assert package_audit["embedded_spreadsheets"] == []
    assert len(package_audit["footer_texts"]) == 3
    assert package_audit["footer_texts"] == ["Page 1", "Page 1", "Page 1"]


@pytest.mark.django_db
def test_xsq_report_uses_the_packaged_aureum_logo_and_omits_legacy_fee_note(report_fixture):
    report, _, tmp_path = report_fixture
    report.fund.short_code = "xsq"
    report.fund.performance_note = (
        "Note: Please refer to the fund offering documents for a detailed fee structure."
    )
    report.fund.save(update_fields=["short_code", "performance_note", "updated_at"])

    _, output = _build_docx(report, tmp_path)
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    asset = Path(reports.AUREUM_INFINITY_LOGO_PATH)

    assert "detailed fee structure" not in text
    with ZipFile(output) as package:
        embedded_hashes = {
            hashlib.sha256(package.read(name)).hexdigest()
            for name in package.namelist()
            if name.startswith("word/media/")
        }
    assert hashlib.sha256(asset.read_bytes()).hexdigest() in embedded_hashes


@pytest.mark.django_db
def test_monthly_docx_uses_monthly_period_and_quarterly_performance_table(report_fixture):
    report, _, tmp_path = report_fixture
    report.report_type = QuarterlyReport.ReportType.MONTHLY
    report.report_month = 2
    report.report_date = date(2024, 2, 29)
    report.commentary_date = report.report_date
    report.save(
        update_fields=[
            "report_type",
            "report_month",
            "report_date",
            "commentary_date",
            "updated_at",
        ]
    )

    snapshot, output = _build_docx(report, tmp_path)
    document = Document(output)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )

    assert snapshot["identity"]["report_type"] == "MONTHLY"
    assert snapshot["identity"]["period_label"] == "2024 年 2 月月報"
    assert snapshot["calculation"]["report_type"] == "MONTHLY"
    assert "2024 年 2 月月報" in paragraph_text
    assert "基金表現（季度淨回報）" in paragraph_text
    assert "基金表現（每月淨回報）" not in paragraph_text
    assert "年度" in table_text
    assert "第一季" in table_text
    assert "年初至今" in table_text
    assert "3.00%" in table_text
    assert all("Version" not in footer.text for footer in document.sections[0].footer.paragraphs)
    package_audit = audit_docx_package(output)
    assert package_audit["valid"] is True
    assert package_audit["external_excel_relationships"] == []


@pytest.mark.django_db
def test_monthly_preview_uses_quarterly_performance_table(report_fixture, client):
    report, user, _ = report_fixture
    report.report_type = QuarterlyReport.ReportType.MONTHLY
    report.report_month = 2
    report.report_date = date(2024, 2, 29)
    report.commentary_date = report.report_date
    report.save(
        update_fields=[
            "report_type",
            "report_month",
            "report_date",
            "commentary_date",
            "updated_at",
        ]
    )
    client.force_login(user)

    response = client.get(reverse("report-preview", args=[report.pk]))

    assert response.status_code == 200
    assert "基金表現－季度淨回報" in response.content.decode()
    assert "基金表現－每月淨回報" not in response.content.decode()
    assert "第一季" in response.content.decode()
    assert "年初至今" in response.content.decode()


@pytest.mark.django_db
def test_monthly_generation_uses_custom_layout_when_template_is_configured(
    report_fixture, monkeypatch
):
    report, user, _ = report_fixture
    report.report_type = QuarterlyReport.ReportType.MONTHLY
    report.report_month = 3
    report.save(update_fields=["report_type", "report_month", "updated_at"])
    report.fund.custom_docx_template.name = "funds/templates/quarterly-only.docx"
    report.fund.save(update_fields=["custom_docx_template", "updated_at"])
    custom_calls = []

    def build_custom(*args, **kwargs):
        custom_calls.append((args, kwargs))
        snapshot, chart_path, _, output_path = args
        build_builtin_docx(snapshot, chart_path, output_path)

    def fake_convert(docx_path, output_dir):
        pdf = output_dir / f"{docx_path.stem}.pdf"
        pdf.write_bytes(b"%PDF-1.4\n% monthly fallback test\n")
        return pdf

    monkeypatch.setattr(reports, "build_custom_docx", build_custom)
    monkeypatch.setattr(reports, "convert_docx_to_pdf", fake_convert)
    generated = generate_report_files(report, user)

    assert len(custom_calls) == 1
    assert {item.file_type for item in generated} == {"DOCX", "PDF"}
    monthly_docx = next(item.absolute_path for item in generated if item.file_type == "DOCX")
    assert "基金表現（季度淨回報）" in "\n".join(
        paragraph.text for paragraph in Document(monthly_docx).paragraphs
    )


@pytest.mark.django_db
def test_builtin_docx_uses_simplified_chinese_system_copy(report_fixture):
    report, _, tmp_path = report_fixture
    report.report_language = QuarterlyReport.ReportLanguage.SIMPLIFIED_CHINESE
    report.save(update_fields=["report_language", "updated_at"])

    snapshot, output = _build_docx(report, tmp_path)
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )

    assert snapshot["identity"]["report_language"] == "zh-Hans"
    assert snapshot["identity"]["period_label"] == "2024年第 1 季季报"
    assert "投资目标" in text
    assert "基金表现（季度净回报）" in text
    assert "基金经理评论" in text
    assert "免责声明" in text
    assert "年度" in table_text
    assert "年初至今" in table_text


@pytest.mark.django_db
def test_builtin_docx_sets_the_selected_cjk_font_on_user_commentary(report_fixture):
    report, _, tmp_path = report_fixture
    report.report_language = QuarterlyReport.ReportLanguage.SIMPLIFIED_CHINESE
    report.commentary_markdown = "简体中文评论必须使用 CJK 字型。"
    report.save(update_fields=["report_language", "commentary_markdown", "updated_at"])

    _, output = _build_docx(report, tmp_path)
    document = Document(output)
    paragraph = next(item for item in document.paragraphs if "简体中文评论" in item.text)

    fonts = paragraph.runs[0]._element.rPr.rFonts
    assert fonts.get(qn("w:eastAsia")) == "Noto Sans CJK SC"
    assert fonts.get(qn("w:ascii")) == "Noto Sans CJK SC"


@pytest.mark.django_db
def test_boya_reference_docx_uses_original_fonts_monthly_table_and_embedded_chart(report_fixture):
    report, _, tmp_path = report_fixture
    report.fund.short_code = "boya"
    report.fund.save(update_fields=["short_code", "updated_at"])
    snapshot = build_current_snapshot(report)
    chart = tmp_path / "boya-chart.png"
    output = tmp_path / "boya.docx"
    generate_nav_chart(snapshot, chart)

    build_boya_reference_docx(snapshot, chart, output)

    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    fonts = document.styles["Normal"]._element.rPr.rFonts
    monthly_header = [cell.text for cell in document.tables[0].rows[0].cells]
    assert "Fund Performance (Net Monthly Returns)" in text
    assert "Fund Performance (Graph)" in text
    assert "Performance Matrix" in text
    matrix_labels = [cell.text for row in document.tables[1].rows for cell in row.cells[::2]]
    assert "Monthly SD (inception)" in matrix_labels
    assert "No of data" in matrix_labels
    assert monthly_header == [
        "Year",
        "Jan",
        "Feb",
        "Mar",
        "Apr",
        "May",
        "Jun",
        "Jul",
        "Aug",
        "Sep",
        "Oct",
        "Nov",
        "Dec",
        "YTD",
    ]
    assert fonts.get(qn("w:ascii")) == "Times New Roman"
    assert fonts.get(qn("w:eastAsia")) == "SimSun"
    assert external_excel_relationships(output) == []
    assert audit_docx_package(output)["valid"] is True
    with ZipFile(output) as package:
        assert len([name for name in package.namelist() if name.startswith("word/media/")]) == 1


def test_boya_performance_matrix_excludes_the_inception_nav_baseline():
    values = [
        "100",
        "99.71759987515605",
        "99.79307677902622",
        "100.84251040799334",
        "92.38147210657785",
        "94.53101207327228",
        "91.38499084096586",
        "88.92",
    ]
    months = [
        date(2025, 11, 30),
        date(2025, 12, 31),
        date(2026, 1, 31),
        date(2026, 2, 28),
        date(2026, 3, 31),
        date(2026, 4, 30),
        date(2026, 5, 31),
        date(2026, 6, 30),
    ]
    points = [
        NavPoint(months[0], date(2025, 11, 24), Decimal(values[0])),
        *[
            NavPoint(month, month, Decimal(value))
            for month, value in zip(months[1:], values[1:], strict=True)
        ],
    ]
    calculation = calculate_performance(
        points=points,
        inception_nav=Decimal("100"),
        inception_date=date(2025, 11, 24),
        report_end=date(2026, 6, 30),
        annual_rfr_decimal=Decimal("0.0423"),
    )
    snapshot = {
        "share_class": {"inception_date": "2025-11-24"},
        "calculation": calculation,
    }
    document = Document()

    reports._add_boya_statistics(document, snapshot)

    values_by_label = {row.cells[0].text: row.cells[1].text for row in document.tables[0].rows}
    assert values_by_label == {
        "Inception Date": "11/24/2025",
        "End Date": "6/30/2026",
        "ITD": "-11.08%",
        "YTD": "-10.83%",
        "Monthly SD (inception)": "3.59%",
        "Monthly SD (12 months)": "3.59%",
        "Days": "218",
        "Annualized Return": "-18.23%",
        "Max Monthly Gain": "2.33%",
        "Max Monthly Loss": "-8.39%",
        "Annualized SD": "12.43%",
        "Trailing 12 Months SD": "12.43%",
        "Rf rate": "4.23%",
        "Sharpe": "-1.807",
        "Max Drawdown": "-11.82%",
        "No of data": "7",
        "% Positive Months": "43%",
        "% Negative Months": "57%",
    }


def test_pdf_conversion_replaces_an_existing_pdf(monkeypatch, settings, tmp_path):
    docx = tmp_path / "report.docx"
    output_dir = tmp_path / "output"
    pdf = output_dir / "report.pdf"
    docx.write_bytes(b"new docx")
    output_dir.mkdir()
    pdf.write_bytes(b"stale pdf")

    monkeypatch.setattr(reports.shutil, "which", lambda _: "soffice")

    def fake_run(*args, **kwargs):
        assert not pdf.exists()
        pdf.write_bytes(b"new pdf")
        return SimpleNamespace(returncode=0, stdout="", stderr="")

    monkeypatch.setattr(reports.subprocess, "run", fake_run)

    assert reports.convert_docx_to_pdf(docx, output_dir) == pdf
    assert pdf.read_bytes() == b"new pdf"


@pytest.mark.django_db
def test_long_commentary_flows_and_snapshot_survives_fund_edit(report_fixture):
    report, _, tmp_path = report_fixture
    report.commentary_markdown = "\n\n".join(
        f"Paragraph {index}: " + ("Long-form portfolio commentary. " * 20) for index in range(1, 35)
    )
    report.save()
    snapshot, output = _build_docx(report, tmp_path)
    assert output.stat().st_size > 10_000
    assert "Paragraph 34" in "\n".join(item.text for item in Document(output).paragraphs)
    report.snapshot = snapshot
    report.save(update_fields=["snapshot", "updated_at"])
    captured = report.snapshot
    report.fund.display_name = "Renamed after capture"
    report.fund.disclaimer_override = "Changed later"
    report.fund.save()
    report.refresh_from_db()
    assert report.snapshot == captured
    assert report.snapshot["identity"]["fund_display_name"] == "Example Fund"


@pytest.mark.django_db
def test_generation_stores_hashes_and_failure_state_persists(report_fixture, monkeypatch):
    report, user, tmp_path = report_fixture

    def fake_convert(docx_path, output_dir):
        pdf = output_dir / f"{docx_path.stem}.pdf"
        pdf.write_bytes(b"%PDF-1.4\n% deterministic test\n")
        return pdf

    monkeypatch.setattr(reports, "convert_docx_to_pdf", fake_convert)
    generated = generate_report_files(report, user)
    report.refresh_from_db()
    assert report.status == QuarterlyReport.Status.READY
    assert {item.file_type for item in generated} == {"DOCX", "PDF"}
    assert all(item.sha256 == sha256_file(item.absolute_path) for item in generated)

    def broken_convert(docx_path, output_dir):
        raise ReportGenerationError("converter deliberately unavailable")

    monkeypatch.setattr(reports, "convert_docx_to_pdf", broken_convert)
    with pytest.raises(ReportGenerationError, match="deliberately unavailable"):
        generate_report_files(report, user)
    report.refresh_from_db()
    assert report.status == QuarterlyReport.Status.GENERATION_FAILED
    assert "deliberately unavailable" in report.generation_error
    assert report.files.filter(file_type=GeneratedFile.FileType.DOCX).exists()


@pytest.mark.django_db
def test_finalized_report_is_immutable_and_nav_edit_marks_it_stale(report_fixture, monkeypatch):
    report, user, _ = report_fixture

    def fake_convert(docx_path, output_dir):
        pdf = output_dir / f"{docx_path.stem}.pdf"
        pdf.write_bytes(b"%PDF-1.4\n% final fixture\n")
        return pdf

    monkeypatch.setattr(reports, "convert_docx_to_pdf", fake_convert)
    generate_report_files(report, user)
    finalize_report(report, user)
    assert report.status == QuarterlyReport.Status.FINAL
    report.commentary_markdown = "Attempted mutation"
    with pytest.raises(ValidationError, match="不可修改"):
        report.save()
    report.refresh_from_db()
    report.status = QuarterlyReport.Status.DRAFT
    with pytest.raises(ValidationError, match="狀態只可"):
        report.save()
    generated = report.files.first()
    generated.sha256 = "0" * 64
    with pytest.raises(ValidationError, match="檔案不可修改"):
        generated.save()
    rfr = report.rfr_snapshot
    rfr.override_reason = "Attempted mutation"
    with pytest.raises(ValidationError, match="無風險利率快照不可修改"):
        rfr.save()
    record = report.share_class.nav_records.get(valuation_month=date(2024, 2, 29))
    assert mark_affected_reports_stale(record, user, "NAV correction") == 1
    report.refresh_from_db()
    assert report.status == QuarterlyReport.Status.STALE


@pytest.mark.django_db
def test_nav_change_resets_mutable_report_snapshot_and_downloads(report_fixture, monkeypatch):
    report, user, _ = report_fixture

    def fake_convert(docx_path, output_dir):
        pdf = output_dir / f"{docx_path.stem}.pdf"
        pdf.write_bytes(b"%PDF-1.4\n% mutable fixture\n")
        return pdf

    monkeypatch.setattr(reports, "convert_docx_to_pdf", fake_convert)
    generate_report_files(report, user)
    report.refresh_from_db()
    assert report.status == QuarterlyReport.Status.READY
    assert report.snapshot
    assert report.files.count() == 2

    record = report.share_class.nav_records.get(valuation_month=date(2024, 2, 29))
    assert reset_affected_mutable_reports(record, user, "NAV correction") == 1
    report.refresh_from_db()
    assert report.status == QuarterlyReport.Status.DRAFT
    assert report.snapshot == {}
    assert report.generation_error == ""
    assert not report.files.exists()


@pytest.mark.django_db
def test_reportable_fund_change_marks_final_report_stale(report_fixture, monkeypatch):
    report, user, _ = report_fixture

    def fake_convert(docx_path, output_dir):
        pdf = output_dir / f"{docx_path.stem}.pdf"
        pdf.write_bytes(b"%PDF-1.4\n% fund-change fixture\n")
        return pdf

    monkeypatch.setattr(reports, "convert_docx_to_pdf", fake_convert)
    generate_report_files(report, user)
    finalize_report(report, user)
    report.fund.display_name = "Updated display name"
    report.fund.save(update_fields=["display_name", "updated_at"])
    assert mark_fund_reports_stale(report.fund, user, "Fund display name changed") == 1
    report.refresh_from_db()
    assert report.status == QuarterlyReport.Status.STALE
    assert report.snapshot["identity"]["fund_display_name"] == "Example Fund"


@pytest.mark.django_db
def test_reportable_share_class_change_marks_final_report_stale(report_fixture, monkeypatch):
    report, user, _ = report_fixture

    def fake_convert(docx_path, output_dir):
        pdf = output_dir / f"{docx_path.stem}.pdf"
        pdf.write_bytes(b"%PDF-1.4\n% share-change fixture\n")
        return pdf

    monkeypatch.setattr(reports, "convert_docx_to_pdf", fake_convert)
    generate_report_files(report, user)
    finalize_report(report, user)
    report.share_class.name = "Class A Updated"
    report.share_class.save(update_fields=["name", "updated_at"])
    assert mark_share_class_reports_stale(report.share_class, user, "Share-class name changed") == 1
    report.refresh_from_db()
    assert report.status == QuarterlyReport.Status.STALE
    assert report.snapshot["identity"]["share_class_name"] == "Class A"


@pytest.mark.django_db
def test_organization_change_marks_final_report_stale(report_fixture, monkeypatch):
    report, user, _ = report_fixture

    def fake_convert(docx_path, output_dir):
        pdf = output_dir / f"{docx_path.stem}.pdf"
        pdf.write_bytes(b"%PDF-1.4\n% organization-change fixture\n")
        return pdf

    monkeypatch.setattr(reports, "convert_docx_to_pdf", fake_convert)
    generate_report_files(report, user)
    finalize_report(report, user)
    organization = report.fund.resolved()
    settings = OrganizationSettings.load()
    settings.percentage_decimal_places = 3
    settings.save(update_fields=["percentage_decimal_places", "updated_at"])
    assert mark_organization_reports_stale(settings, user, "Presentation precision changed") == 1
    report.refresh_from_db()
    assert report.status == QuarterlyReport.Status.STALE
    assert organization["percentage_decimal_places"] == 2


@pytest.mark.integration
@pytest.mark.django_db
@pytest.mark.skipif(shutil.which("soffice") is None, reason="LibreOffice is not installed")
def test_real_libreoffice_pdf_conversion(report_fixture):
    report, user, _ = report_fixture
    files = generate_report_files(report, user)
    pdf = next(item for item in files if item.file_type == GeneratedFile.FileType.PDF)
    assert pdf.absolute_path.read_bytes().startswith(b"%PDF")
    assert pdf.size > 1_000
